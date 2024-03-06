import os
import subprocess

from docx import Document
import platform

def get_word_document_text(file):
    """
    Reads a Word document and returns the text content
    """
    doc = Document(file)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def get_word_document_headers(file):
    """
    Reads a Word document and returns the headers
    """
    doc = Document(file)
    headers = []
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
            headers.append(para.text)
    return headers

def doc_to_docx(doc_path):
    if platform.system() == 'Windows':
        return doc_to_docx_windows(doc_path)
    elif platform.system() in ['Linux', 'Darwin']:  # 'Darwin' for macOS
        return doc_to_docx_linux(doc_path)
    else:
        print("Unsupported operating system.")
        return None

def doc_to_docx_windows(doc_path):
    import comtypes.client

    if doc_path.endswith('.docx'):
        return

    # Initialize the Word application
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False

    # create docx from doc
    docx_path = doc_path + "x"
    # Open the .doc file
    doc = word.Documents.Open(doc_path)

    # Save it as a .docx file
    doc.SaveAs(docx_path, FileFormat=16)  # 16 represents the wdFormatXMLDocument constant

    # Close Word
    doc.Close()
    word.Quit()

    return docx_path

os.environ["PATH"] += os.pathsep + "/usr/bin"
def doc_to_docx_linux(doc_path):
    # Check if the file is already a .docx file
    if doc_path.endswith('.docx'):
        return doc_path

    # Create the .docx file path
    docx_path = os.path.splitext(doc_path)[0] + '.docx'

    # Convert the file using unoconv
    try:
        subprocess.run(['unoconv', '-f', 'docx', '-o', docx_path, doc_path], check=True)
    except subprocess.CalledProcessError as e:
        print(f"Error during conversion: {e}")
        return None

    return docx_path